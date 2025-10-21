#Python标准库
import os
from typing import List, Dict, Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor
import pytz
from datetime import datetime, timezone, timedelta

#第三方库
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import  A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 注册中文字体（关键步骤）
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 自定义库
from services.service_models import Meeting, Transcription


TABLE_STYLE = 'Table Grid'
DATETIME_CHINESE_SIMPLE = "%Y年%m月%d日 %H:%M"

class DocumentService(object):
    def __init__(self)-> None:
        self.executor = ThreadPoolExecutor(max_workers=4)
        self.output_dir = "static/documents"
        self.east8_tz = timezone(timedelta(hours=8))
        os.makedirs(self.output_dir, exist_ok=True)

        # 表格样式配置（统一管理，修改时仅需改一处）
        self.table_config = {
            "align": WD_TABLE_ALIGNMENT.CENTER,
            "header_font_size": Pt(11),
            "content_font_size": Pt(10),
            "col_widths": [Inches(1.2), Inches(3.5)]  # 基本信息表格：左列1.2英寸，右列3.5英寸
        }
        self.default_font = "微软雅黑"  # 中文适配字体
        # Register Chinese font for PDF generation
        try:
            # Try to use a Chinese font (you might need to install one)
            font_path = r"D:/PycharmProject/DejaVuSans.ttf"
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('Chinese', font_path))

        except (IOError, pdfmetrics.FontError) as e:
            print(f"字体注册失败({font_path}): {str(e)}")

    def _convert_to_east8_time(self, dt: datetime) -> datetime:
        """将时间转换为东八区时间（使用pytz）"""
        if dt.tzinfo is None:
            # 如果时间没有时区信息，假定为UTC时间
            dt = dt.replace(tzinfo=timezone.utc)

        # 转换为东八区
        east8_tz = pytz.timezone('Asia/Shanghai')
        return dt.astimezone(east8_tz)

    async def generate_notification(self, meeting: Meeting):
        """Generate meeting notification document in both Word and PDF formats"""
        # Generate Word document
        print("开始打印通知书")
        word_path = await self._generate_notification_word(meeting)
        # Generate PDF document
        pdf_path = await self._generate_notification_pdf(meeting)
        return {"word": word_path, "pdf": pdf_path}

    async def _generate_notification_word(self, meeting: Meeting) -> str:
        """Generate Word format meeting notification"""
        doc = Document()
        # Add title
        title = doc.add_heading('会议通知', 0)
        # Center alignment
        title.alignment = 1
        # Add meeting details
        doc.add_heading('会议信息', level=1)
        details_table = doc.add_table(rows=6, cols=2)
        details_table.style = TABLE_STYLE
        # Fill table with meeting information
        cells = details_table.rows[0].cells
        cells[0].text = '会议主题'
        cells[1].text = meeting.title
        cells = details_table.rows[1].cells
        cells[0].text = '会议时间'
        cells[1].text = meeting.date_time.strftime(DATETIME_CHINESE_SIMPLE)
        cells = details_table.rows[2].cells
        cells[0].text = '会议地点'
        cells[1].text = meeting.location or '待定'
        cells = details_table.rows[3].cells
        cells[0].text = '预计时长'
        cells[1].text = f'{meeting.duration_minutes}分钟'
        cells = details_table.rows[4].cells
        cells[0].text = '会议描述'
        cells[1].text = meeting.description or '无'
        cells = details_table.rows[5].cells
        cells[0].text = '会议议程'
        cells[1].text = meeting.agenda or '待补充'
        # Add participants section
        if meeting.participants:
            doc.add_heading('参会人员', level=1)
            participants_table = doc.add_table(rows=1, cols=3)
            participants_table.style = TABLE_STYLE
            # Header row
            header_cells = participants_table.rows[0].cells
            header_cells[0].text = '姓名'
            header_cells[1].text = '邮箱'
            header_cells[2].text = '角色'
            # Add participant rows
            for participant in meeting.participants:
                row_cells = participants_table.add_row().cells
                row_cells[0].text = participant.name
                row_cells[1].text = participant.email
                row_cells[2].text = self._translate_role(participant.user_role)
        # Add footer
        doc.add_paragraph('')
        footer = doc.add_paragraph('请及时确认参会状态，如有冲突请提前告知。')
        footer.add_run(f'\n\n生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        # Save document
        filename = f"meeting_notification_{meeting.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath


    async def _generate_notification_pdf(self, meeting: Meeting) -> str:
        """Generate PDF format meeting notification asynchronously"""
        # 在线程池中执行CPU密集型的PDF生成操作
        loop = asyncio.get_event_loop()
        try:
            filepath = await loop.run_in_executor(
                self.executor,
                self._generate_pdf_sync,
                meeting
            )
            return filepath
        except Exception as e:
            print(f"Failed to generate PDF notification: {e}")
            raise

    def _generate_pdf_sync(self, meeting: Meeting) -> str:
        """Synchronous PDF generation (CPU-intensive)"""
        filename = f"meeting_notification_{meeting.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)

        # 确保系统中存在中文字体文件，或者提供字体文件路径
        # 这里使用常见的 SimHei（黑体）作为示例，你也可以使用其他中文字体
        TTFONT = 'SimHei.ttf'
        try:
            pdfmetrics.registerFont(TTFont('ChineseFont', TTFONT))
        except (IOError, pdfmetrics.FontError) as e:
            # 定义常见中文字体列表，按优先级排序
            chinese_fonts = [
                (TTFONT, '黑体'),
                ('msyh.ttf', '微软雅黑'),
                ('simsun.ttc', '宋体')
            ]

            font_registered = False
            last_error = str(e)

            for font_file, font_name in chinese_fonts:
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_file))
                    font_registered = True
                    break
                except (IOError, pdfmetrics.FontError) as font_error:
                    last_error = f"{font_name}({font_file}): {str(font_error)}"
                    continue

            if not font_registered:
                print(f"警告：中文字体注册失败，尝试了以下字体:\n{last_error}\n中文可能显示为乱码")

        doc = SimpleDocTemplate(filepath, pagesize=A4)

        styles = getSampleStyleSheet()

        # 创建自定义样式 - 全部使用中文字体
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center
            fontName='ChineseFont'  # 使用中文字体
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            fontName='ChineseFont'  # 使用中文字体
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            fontName='ChineseFont'  # 使用中文字体
        )

        story = []

        # 标题
        story.append(Paragraph("会议通知", title_style))
        story.append(Spacer(1, 20))

        # 会议详情
        story.append(Paragraph("会议信息", heading_style))

        meeting_data = [
            ['会议主题', meeting.title],
            ['会议时间', meeting.date_time.strftime(DATETIME_CHINESE_SIMPLE)],
            ['会议地点', meeting.location or '待定'],
            ['预计时长', f'{meeting.duration_minutes}分钟'],
            ['会议描述', meeting.description or '无'],
            ['会议议程', meeting.agenda or '待补充']
        ]

        meeting_table = Table(meeting_data, colWidths=[2 * inch, 4 * inch])
        meeting_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'ChineseFont'),  # 表格也使用中文字体
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(meeting_table)
        story.append(Spacer(1, 20))

        # 参会人员
        if meeting.participants:
            story.append(Paragraph("参会人员", heading_style))

            participant_data = [['姓名', '邮箱', '角色']]
            for participant in meeting.participants:
                participant_data.append([
                    participant.name,
                    participant.email,
                    self._translate_role(participant.user_role)
                ])

            participants_table = Table(participant_data, colWidths=[2 * inch, 2.5 * inch, 1.5 * inch])
            participants_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'ChineseFont'),  # 表格使用中文字体
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(participants_table)

        # 页脚
        story.append(Spacer(1, 30))
        story.append(Paragraph("请及时确认参会状态，如有冲突请提前告知。", normal_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"生成时间：{datetime.now().strftime(DATETIME_CHINESE_SIMPLE)}", normal_style))

        doc.build(story)
        return filepath

    def _translate_role(self, role: str) -> str:
        """Translate role to Chinese"""
        role_translations = {
            'organizer': '组织者',
            'participant': '参会者',
            'presenter': '演讲者',
            'guest': '嘉宾'
        }
        return role_translations.get(role, '参会者')

    async def generate_minutes(self, meeting: Meeting, transcriptions: list[Transcription]) ->  dict[str, str]:
        """Generate meeting minutes document"""
        # Generate PDF document
        pdf_path = await self._generate_minutes_pdf(meeting, transcriptions)

        # Generate Word document
        word_path = await self._generate_minutes_word(meeting, transcriptions)

        return {"word": word_path, "pdf": pdf_path}

    async def _generate_minutes_pdf(self, meeting: Meeting, transcriptions: list[Transcription]) -> str:
        """Generate PDF format meeting minutes"""
        filename = f"meeting_minutes_{meeting.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)

        # 注册中文字体
        chinese_font_name = self._register_chinese_font()

        # 创建文档和样式
        doc, styles = self._create_document_and_styles(filepath, chinese_font_name)

        # 构建文档内容
        story = self._build_pdf_content(meeting, transcriptions, styles, chinese_font_name)

        # 生成PDF
        doc.build(story)
        return filepath

    def _create_document_and_styles(self, filepath: str, chinese_font_name: str) -> tuple:
        """创建PDF文档和样式"""
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        styles = getSampleStyleSheet()

        # 定义样式
        custom_styles = {
            'title_style': ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1,
                fontName=chinese_font_name
            ),
            'heading_style': ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                fontName=chinese_font_name
            ),
            'normal_style': ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                fontName=chinese_font_name
            )
        }

        return doc, custom_styles

    def _build_pdf_content(self, meeting: Meeting, transcriptions: list[Transcription],
                           styles: dict, chinese_font_name: str) -> list:
        """构建PDF内容"""
        story = []

        # 添加标题
        self._add_title_section(story, styles['title_style'])

        # 添加会议基本信息
        self._add_meeting_info_section(story, meeting, styles['heading_style'],
                                       styles['normal_style'], chinese_font_name)

        # 添加会议内容
        self._add_transcriptions_section(story, transcriptions, styles['heading_style'],
                                         styles['normal_style'])

        # 添加行动项汇总
        self._add_action_items_section(story, transcriptions, styles['heading_style'],
                                       styles['normal_style'])

        # 添加决议汇总
        self._add_decisions_section(story, transcriptions, styles['heading_style'],
                                    styles['normal_style'])

        # 添加页脚
        self._add_footer_section(story, styles['normal_style'])

        return story

    def _add_title_section(self, story: list, title_style: ParagraphStyle) -> None:
        """添加标题部分"""
        story.append(Paragraph("会议纪要", title_style))
        story.append(Spacer(1, 20))

    def _add_meeting_info_section(self, story: list, meeting: Meeting,
                                  heading_style: ParagraphStyle, normal_style: ParagraphStyle,
                                  chinese_font_name: str) -> None:
        """添加会议信息部分"""
        story.append(Paragraph("会议基本信息", heading_style))

        meeting_data = [
            ['会议主题', meeting.title],
            ['会议时间', meeting.date_time.strftime(DATETIME_CHINESE_SIMPLE)],
            ['会议地点', meeting.location or '线上会议'],
            ['参会人数', str(len(meeting.participants))]
        ]

        meeting_table = Table(meeting_data, colWidths=[2 * inch, 4 * inch])
        meeting_table.setStyle(self._get_table_style(chinese_font_name))

        story.append(meeting_table)
        story.append(Spacer(1, 20))

    def _get_table_style(self, chinese_font_name: str) -> TableStyle:
        """获取表格样式"""
        return TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), chinese_font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])

    def _add_transcriptions_section(self, story: list, transcriptions: list[Transcription],
                                    heading_style: ParagraphStyle, normal_style: ParagraphStyle) -> None:
        """添加会议内容部分"""
        if not transcriptions:
            return

        story.append(Paragraph("会议内容", heading_style))

        for transcription in transcriptions:
            content = self._format_transcription_content(transcription)
            story.append(Paragraph(content, normal_style))
            story.append(Spacer(1, 6))

    def _format_transcription_content(self, transcription: Transcription) -> str:
        """格式化转录内容"""
        timestamp = self._convert_to_east8_time(transcription.timestamp).strftime('%H:%M:%S')
        speaker = transcription.speaker_name or transcription.speaker_id
        content = f"[{timestamp}] {speaker}: {transcription.text}"

        if transcription.is_action_item:
            content += " [行动项]"
        if transcription.is_decision:
            content += " [决议]"

        return content

    def _add_action_items_section(self, story: list, transcriptions: list[Transcription],
                                  heading_style: ParagraphStyle, normal_style: ParagraphStyle) -> None:
        """添加行动项部分"""
        action_items = [t for t in transcriptions if t.is_action_item]
        if not action_items:
            return

        story.append(Spacer(1, 20))
        story.append(Paragraph("行动项汇总", heading_style))

        for i, item in enumerate(action_items, 1):
            story.append(Paragraph(f"{i}. {item.text}", normal_style))
            story.append(Spacer(1, 6))

    def _add_decisions_section(self, story: list, transcriptions: list[Transcription],
                               heading_style: ParagraphStyle, normal_style: ParagraphStyle) -> None:
        """添加决议部分"""
        decisions = [t for t in transcriptions if t.is_decision]
        if not decisions:
            return

        story.append(Spacer(1, 20))
        story.append(Paragraph("重要决议", heading_style))

        for i, decision in enumerate(decisions, 1):
            story.append(Paragraph(f"{i}. {decision.text}", normal_style))
            story.append(Spacer(1, 6))

    def _add_footer_section(self, story: list, normal_style: ParagraphStyle) -> None:
        """添加页脚部分"""
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            f"会议纪要生成时间：{datetime.now().strftime(DATETIME_CHINESE_SIMPLE)}",
            normal_style
        ))


    def _register_chinese_font(self) -> str:
        """注册中文字体并返回字体名称"""
        from reportlab.pdfbase import pdfmetrics

        # 按优先级尝试不同的字体注册方法
        font_attempts = [
            self._try_ttf_fonts,
            self._try_cid_font
        ]

        for attempt in font_attempts:
            font_name = attempt(pdfmetrics)
            if font_name:
                return font_name

        print("警告：未找到中文字体，中文可能显示为乱码")
        return 'Helvetica'

    def _try_ttf_fonts(self, pdfmetrics: str) -> Optional[str]:
        """尝试注册TTF字体"""
        from reportlab.pdfbase.ttfonts import TTFont

        font_candidates = [
            ('SimHei', 'SimHei.ttf'),
            ('MicrosoftYaHei', 'msyh.ttf'),
            ('SimSun', 'simsun.ttc'),
            ('STSong', 'STSONG.TTF'),
        ]

        system_font_paths = [
            '/usr/share/fonts/truetype/',
            '/usr/share/fonts/truetype/msttcorefonts/',
            'C:/Windows/Fonts/',
            '/System/Library/Fonts/',
            '/Library/Fonts/'
        ]

        for font_name, font_file in font_candidates:
            # 尝试直接注册
            if self._register_font(pdfmetrics, TTFont, font_name, font_file):
                return font_name

            # 尝试系统路径注册
            system_font = self._find_in_system_paths(
                pdfmetrics, TTFont, font_name, font_file, system_font_paths
            )
            if system_font:
                return system_font

        return None

    @staticmethod
    def _try_cid_font(pdfmetrics: str) -> Optional[str]:
        """尝试注册CID字体"""
        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            return 'STSong-Light'
        except Exception:
            return None


    def _register_font(self, pdfmetrics: str, font_class: str, font_name: str, font_path: str) -> bool:
        """注册单个字体"""
        try:
            pdfmetrics.registerFont(font_class(font_name, font_path))
            return True
        except Exception:
            return False

    def _find_in_system_paths(self, pdfmetrics:str, font_class: str, font_name:str, font_file:str, system_paths: list[str]) -> Optional[str]:
        """在系统路径中查找字体"""
        for path in system_paths:
            full_path = os.path.join(path, font_file)
            if os.path.exists(full_path) and self._register_font(pdfmetrics, font_class, font_name, full_path):
                return font_name
        return None


    def _translate_role(self, role: str) -> str:
        """Translate role to Chinese"""
        role_map = {
            'organizer': '组织者',
            'participant': '参与者',
            'presenter': '主讲人'
        }
        return role_map.get(role, role)


    async def _generate_minutes_word(self, meeting: Meeting, transcriptions: list[Transcription]) -> str:
        """Generate Word format meeting minutes"""
        doc = Document()

        # 构建文档的各个部分
        self._add_document_title(doc)
        self._add_meeting_details(doc, meeting)
        self._add_transcription_content(doc, transcriptions)
        self._add_action_items_summary(doc, transcriptions)
        self._add_decisions_summary(doc, transcriptions)
        self._add_document_footer(doc)

        # 保存文档
        return self._save_document(doc, meeting)


    def _add_document_title(self, doc: Document) -> None:
        """添加文档标题"""
        title = doc.add_heading('会议纪要', 0)
        title.alignment = 1  # 居中对齐


    def _add_meeting_details(self, doc: Document, meeting: Meeting) -> None:
        """添加会议基本信息"""
        doc.add_heading('会议基本信息', level=1)

        details_table = doc.add_table(rows=4, cols=2)
        details_table.style = TABLE_STYLE

        # 填充会议信息表格
        details_data = [
            ('会议主题', meeting.title),
            ('会议时间', meeting.date_time.strftime(DATETIME_CHINESE_SIMPLE)),
            ('会议地点', meeting.location or '线上会议'),
            ('参会人数', str(len(meeting.participants)))
        ]

        for i, (label, value) in enumerate(details_data):
            cells = details_table.rows[i].cells
            cells[0].text = label
            cells[1].text = value

    def _add_transcription_content(self, doc: Document, transcriptions: list[Transcription]) -> None:
        """添加转录内容"""
        if not transcriptions:
            return

        doc.add_heading('会议内容', level=1)

        current_speaker = None
        for transcription in transcriptions:
            if transcription.speaker_name != current_speaker:
                current_speaker = transcription.speaker_name or transcription.speaker_id
                doc.add_heading(f'{current_speaker}:', level=3)

            self._add_transcription_paragraph(doc, transcription)

    def _add_transcription_paragraph(self, doc: Document, transcription: Transcription) -> None:
        """添加单个转录段落"""

        timestamp = self._convert_to_east8_time(transcription.timestamp).strftime('%H:%M:%S')
        paragraph = doc.add_paragraph(f'[{timestamp}] {transcription.text}')

        # 高亮行动项和决议
        if transcription.is_action_item:
            paragraph.add_run(' [行动项]').bold = True
        if transcription.is_decision:
            paragraph.add_run(' [决议]').bold = True


    def _add_action_items_summary(self, doc: Document, transcriptions: list[Transcription]) -> None:
        """添加行动项汇总"""
        action_items = [t for t in transcriptions if t.is_action_item]
        if not action_items:
            return

        doc.add_heading('行动项汇总', level=1)
        for i, item in enumerate(action_items, 1):
            doc.add_paragraph(f'{i}. {item.text}', style='List Number')

    def _add_decisions_summary(self, doc: Document, transcriptions: list[Transcription]) -> None:
        """添加决议汇总"""
        decisions = [t for t in transcriptions if t.is_decision]
        if not decisions:
            return

        doc.add_heading('重要决议', level=1)
        for i, decision in enumerate(decisions, 1):
            doc.add_paragraph(f'{i}. {decision.text}', style='List Number')

    def _add_document_footer(self, doc: Document) -> None:
        """添加文档页脚"""
        doc.add_paragraph('')
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")
        doc.add_paragraph(f'会议纪要生成时间：{current_time}')

    def _save_document(self, doc: Document, meeting: Meeting) -> str:
        """保存文档并返回文件路径"""
        filename = f"meeting_minutes_{meeting.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath




